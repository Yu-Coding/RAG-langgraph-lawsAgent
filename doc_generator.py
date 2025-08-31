from docx import Document
import os
#print("当前工作目录:", os.getcwd())


def generate_contract_doc(landlord, tenant, location, amount, duration, filename="租赁合同.docx"):
    doc = Document()
    doc.add_heading("房屋租赁合同", level=0)

    doc.add_paragraph(f"甲方（房东）：{landlord}")
    doc.add_paragraph(f"乙方（租客）：{tenant}")
    doc.add_paragraph(f"租赁地址：{location}")
    doc.add_paragraph(f"租赁期限：{duration}")
    doc.add_paragraph(f"租金金额：{amount}")
    doc.add_paragraph("双方根据《中华人民共和国民法典》及相关法律法规的规定，本着平等自愿、公平诚信的原则，订立本合同，并共同遵守。")

    doc.add_paragraph("【租赁条款摘要】：")
    doc.add_paragraph("1. 房东应保证房屋产权清晰，无纠纷。")
    doc.add_paragraph("2. 租客应按时支付租金，合理使用房屋。")
    doc.add_paragraph("3. 如需提前解除合同，须提前30天书面通知对方。")
    
    filename = os.path.join(os.path.dirname(__file__), "租赁合同.docx")
    doc.save(filename)
    return filename
